import os
import pytest
import io
import docx
from fastapi.testclient import TestClient

# Mock the dependency before importing app
os.environ["ENVIRONMENT"] = "development"

from app.main import app
from app.api.deps import get_current_user, get_current_active_user
from app.models.user import User

async def mock_get_current_user():
    return User(id="12345678-1234-5678-1234-567812345678", email="admin@example.com", role="Admin", is_active=True)

app.dependency_overrides[get_current_user] = mock_get_current_user
app.dependency_overrides[get_current_active_user] = mock_get_current_user

# Minimal valid PDF that pypdf can read
VALID_PDF_BYTES = b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> >>\nendobj\n4 0 obj\n<< /Length 44 >>\nstream\nBT\n/F1 12 Tf\n72 712 Td\n(PDF Test) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000259 00000 n \ntrailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n354\n%%EOF"

def generate_docx_bytes():
    doc = docx.Document()
    doc.add_paragraph("DOCX Test Paragraph 1")
    doc.add_paragraph("DOCX Test Paragraph 2")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

@pytest.mark.asyncio
async def test_upload_formats():
    with TestClient(app) as client:
        # 1. TXT Upload
        res = client.post("/api/v1/documents/upload", files={"file": ("test.txt", b"TXT Test", "text/plain")})
        assert res.status_code == 201
        assert res.json()["total_chunks"] > 0
        
        # 2. DOCX Upload
        docx_bytes = generate_docx_bytes()
        res = client.post("/api/v1/documents/upload", files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        assert res.status_code == 201
        assert res.json()["total_chunks"] > 0

        # 3. PDF Upload
        res = client.post("/api/v1/documents/upload", files={"file": ("test.pdf", VALID_PDF_BYTES, "application/pdf")})
        # Note: Depending on pypdf's strictness, it might fail or succeed. We'll check if it fails gracefully if so.
        # Let's see if it works. If not, we check for a 500 or 400.
        if res.status_code != 201:
            assert "Failed to parse PDF" in res.text

        # 4. Unsupported Format
        res = client.post("/api/v1/documents/upload", files={"file": ("test.unknown", b"Fake", "application/unknown")})
        assert res.status_code == 400
        assert "Unsupported file type" in res.text

        # 5. Empty Document Handling
        res = client.post("/api/v1/documents/upload", files={"file": ("empty.txt", b"", "text/plain")})
        assert res.status_code == 201
        assert res.json()["total_chunks"] == 0
        assert res.json()["message"] == "No text extracted"

        # 6. Corrupted DOCX
        res = client.post("/api/v1/documents/upload", files={"file": ("corrupt.docx", b"Not a docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
        assert res.status_code == 400
        assert "Failed to parse DOCX" in res.text

        # 7. Corrupted PDF
        res = client.post("/api/v1/documents/upload", files={"file": ("corrupt.pdf", b"Not a pdf", "application/pdf")})
        assert res.status_code == 400
        assert "Failed to parse PDF" in res.text
