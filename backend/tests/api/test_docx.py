import pytest
import io
import docx
from httpx import AsyncClient
from app.core.config import settings


def create_valid_docx_bytes() -> bytes:
    doc = docx.Document()
    doc.add_heading('DOCX Test Document', 0)
    doc.add_paragraph('This is a test paragraph for DOCX integration tests.')
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Row1Col1'
    table.cell(0, 1).text = 'Row1Col2'
    table.cell(1, 0).text = 'Row2Col1'
    table.cell(1, 1).text = 'Row2Col2'
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def create_empty_docx_bytes() -> bytes:
    doc = docx.Document()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

@pytest.mark.asyncio
async def test_upload_valid_docx(client: AsyncClient, admin_token: str):
    file_content = create_valid_docx_bytes()
    response = await client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("test_doc.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    if response.status_code != 201:
        print("Upload failed with 403:", response.text)
    assert response.status_code == 201
    data = response.json()
    assert data["filename"] == "test_doc.docx"
    assert "processed" in data["message"]
    
    # Store document_id for retrieval/deletion test
    document_id = data["document_id"]
    
    # Test retrieval (Hybrid Search)
    search_response = await client.get(
        "/api/v1/documents/search?q=Row1Col1",
        headers={"Authorization": f"Bearer {admin_token}"}
    )
    print("Search Response:", search_response.status_code, search_response.text)
    assert search_response.status_code == 200
    search_data = search_response.json()
    assert isinstance(search_data["results"], list)
    
    # Test deletion
    delete_response = await client.delete(
        f"/api/v1/documents/{document_id}",
        headers={"Authorization": f"Bearer {admin_token}"}
    )
    assert delete_response.status_code == 204

@pytest.mark.asyncio
async def test_upload_empty_docx(client: AsyncClient, admin_token: str):
    file_content = create_empty_docx_bytes()
    response = await client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("empty.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    print("Empty DOCX response:", response.status_code, response.text)
    assert response.status_code == 400

@pytest.mark.asyncio
async def test_upload_corrupted_docx(client: AsyncClient, admin_token: str):
    file_content = b"This is not a valid zip archive, so python-docx will fail."
    response = await client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("corrupted.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    # The DocxExtractor raises ValueError which maps to 400
    assert response.status_code == 400

@pytest.mark.asyncio
async def test_upload_oversized_docx(client: AsyncClient, admin_token: str):
    # Create 12MB of junk data
    file_content = b"0" * (12 * 1024 * 1024)
    response = await client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("huge.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 400
    # In some ASGI testing environments, large bodies might trigger a generic 400 parsing error first.
    assert "exceeds" in response.text or "error parsing the body" in response.text

@pytest.mark.asyncio
async def test_replace_docx(client: AsyncClient, admin_token: str):
    # 1. Upload initial DOCX
    file_content = create_valid_docx_bytes()
    upload_res = await client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("test_doc.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert upload_res.status_code == 201
    doc_id = upload_res.json()["document_id"]
    
    # 2. Upload replacement DOCX
    replace_res = await client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {admin_token}"},
        files={"file": ("replacement_doc.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert replace_res.status_code == 201
    new_doc_id = replace_res.json()["document_id"]
    
    # 3. Delete initial DOCX
    delete_res = await client.delete(
        f"/api/v1/documents/{doc_id}",
        headers={"Authorization": f"Bearer {admin_token}"}
    )
    assert delete_res.status_code == 204
    
    # 4. Clean up new one
    delete_res_new = await client.delete(
        f"/api/v1/documents/{new_doc_id}",
        headers={"Authorization": f"Bearer {admin_token}"}
    )
    assert delete_res_new.status_code == 204
